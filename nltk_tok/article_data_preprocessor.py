# -*- coding: utf-8 -*-
"""
Created on Fri Feb  2 16:57:35 2018

@author: kartikaya
"""

import os
import re
import nltk
from nltk.corpus import stopwords
from nltk.stem.porter import PorterStemmer
from nltk.probability import FreqDist
import operator
from article_data_preprocessor_util import word_count, hit_count
import docx
import pandas as pd

real_estate_keywords = ['house', 'propert', 'properti','property','home', 'mortgag', 'mortgage','estate', 'estat']
retail_keywords = ['retail', 'retai', 'shop', 'amazon' , 'sale', 'supermarket' , 'groceri']
energy_keywords = ['oil','commod', 'gas', 'shale', 'shal', 'solar', 'e-car', 'e-vehicle', 'middle east', 'opec', 'shell', 'bp', 'chevron', 'barrel' , 'energy', 'energi']

xcel = pd.read_excel('C:\\Users\\kartikaya\\Python_Data\\Article_Search_MindTrust\\FUNDS\\Funds-sectors.xlsx')
funds = xcel.iloc[:, 0]
tags = xcel.iloc[:, 1]
isin = xcel.iloc[:, 2]
all_files = os.listdir("C:\\Users\\kartikaya\\Python_Data\\Article_Search_MindTrust\\Articles")

file = 'C:\\Users\\kartikaya\\Python_Data\\Article_Search_MindTrust\\Articles\\' + all_files[7]



doc = docx.Document(file)

full_text = []

for para in doc.paragraphs:
    full_text.append(para.text)
    '\n'.join(full_text)
    
article_cleaned = ''    
for text in full_text:
    text = re.sub('[^a-zA-Z]',' ',text)
    article_cleaned = article_cleaned + text
    
article_cleaned = article_cleaned.lower()
article_cleaned = article_cleaned.split()
ps = PorterStemmer()

article_cleaned = [ps.stem(word) for word in article_cleaned if not word in set(stopwords.words('english'))]
article_cleaned = ' '.join(article_cleaned)

"""
words = nltk.tokenize.word_tokenize(article_cleaned)

fdist = FreqDist(words)

fdist_keys = fdist.keys()
fdist_values = fdist.values()
"""


count = word_count(article_cleaned)

sorted_x = sorted(count.items(), key=operator.itemgetter(1), reverse = True)
real_estate_hits = hit_count(sorted_x, real_estate_keywords)
retail_hits  = hit_count(sorted_x , retail_keywords)
energy_hits = hit_count(sorted_x, energy_keywords)

print('reatail hits - '+str(retail_hits))
print('real estate hits - '+str(real_estate_hits))
print('energy hits - '+str(energy_hits))
            



real_estate_keywords = ['house', 'propert', 'properti','property','home', 'mortgag', 'mortgage','estate', 'estat']
retail_keywords = ['retail', 'retai', 'shop', 'amazon' , 'sale', 'supermarket' , 'groceri']
energy_keywords = ['oil','commod', 'gas', 'shale', 'shal', 'solar', 'e-car', 'e-vehicle', 'middle east', 'opec', 'shell', 'bp', 'chevron', 'barrel']

xcel = pd.read_excel('C:\\Users\\kartikaya\\Python_Data\\Article_Search_MindTrust\\FUNDS\\Funds-sectors.xlsx')
funds = xcel.iloc[:, 0]
tags = xcel.iloc[:, 1]
isin = xcel.iloc[:, 2]
all_files = os.listdir("C:\\Users\\kartikaya\\Python_Data\\Article_Search_MindTrust\\Articles")

i = 1
print('Please select a fund')
for fund in funds:
    print(str(i)+" - "+str(fund))
    i = i + 1

fund = input(" ")
fund = fund.strip()
index = int(fund) - 1
string = funds[index]
print('you selected : '+string)
fund_tags = tags[index]
fund_isin = isin[index]

ps = PorterStemmer()
j = 0
article_list = []
for article in all_files:
    file = 'C:\\Users\\kartikaya\\Python_Data\\Article_Search_MindTrust\\Articles\\' + all_files[j]
    j = j + 1
    doc = docx.Document(file)
    full_text = []
    for para in doc.paragraphs:
       full_text.append(para.text)
       '\n'.join(full_text)
    article_cleaned = ''
    for text in full_text:
       text = re.sub('[^a-zA-Z]',' ',text)
       article_cleaned = article_cleaned + text
    article_cleaned = article_cleaned.lower()
    article_cleaned = article_cleaned.split()
    article_cleaned = [ps.stem(word) for word in article_cleaned if not word in set(stopwords.words('english'))]
    article_cleaned = ' '.join(article_cleaned)
    count = word_count(article_cleaned)
    sorted_x = sorted(count.items(), key=operator.itemgetter(1), reverse = True)
    real_estate_hits = hit_count(sorted_x, real_estate_keywords)
    retail_hits  = hit_count(sorted_x , retail_keywords)
    energy_hits = hit_count(sorted_x, energy_keywords)
    if fund_tags == 'energy' and energy_hits > 10:
        article_list.append(article)
    if fund_tags == 'real_estate' and real_estate_hits > 10:
        article_list.append(article)
print('------------------------------------------------------------')
i=1
for article in article_list:
    print(str(i) + '-> '+ article)
    i = i + 1

    
    
                                                                                                                                                                                
                                                                                                                                                    
                                                                                                                            
import os
import re
import nltk
from nltk.corpus import stopwords
from nltk.stem.porter import PorterStemmer
from nltk.probability import FreqDist
import operator
from article_data_preprocessor_util import word_count, hit_count
import docx
import pandas as pd
real_estate_keywords = ['house', 'propert', 'properti','property','home', 'mortgag', 'mortgage','estate', 'estat']
retail_keywords = ['retail', 'retai', 'shop', 'amazon' , 'sale', 'supermarket' , 'groceri']
energy_keywords = ['oil','commod', 'gas', 'shale', 'shal', 'solar', 'e-car', 'e-vehicle', 'middle east', 'opec', 'shell', 'bp', 'chevron', 'barrel','energy', 'energi']

xcel = pd.read_excel('C:\\Users\\kartikaya\\Python_Data\\Article_Search_MindTrust\\FUNDS\\Funds-sectors.xlsx')
funds = xcel.iloc[:, 0]
tags = xcel.iloc[:, 1]
isin = xcel.iloc[:, 2]
all_files = os.listdir("C:\\Users\\kartikaya\\Python_Data\\Article_Search_MindTrust\\Articles")    
while 1 == 1:
    i = 1
    print('------------------------------------------------------------')
    print('------------------------------------------------------------')
    print('Please select a fund')
    for fund in funds:
        print(str(i)+" - "+str(fund))
        i = i + 1
    print('------------------------------------------------------------')
    print('------------------------------------------------------------')
    fund = input(" ")
    fund = fund.strip()
    index = int(fund) - 1
    string = funds[index]
    print('you selected : '+string)
    fund_tags = tags[index]
    fund_isin = isin[index]

    ps = PorterStemmer()
    j = 0
    article_list = []
    for article in all_files:
        file = 'C:\\Users\\kartikaya\\Python_Data\\Article_Search_MindTrust\\Articles\\' + all_files[j]
        j = j + 1
        doc = docx.Document(file)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
            '\n'.join(full_text)
        article_cleaned = ''
        for text in full_text:
            text = re.sub('[^a-zA-Z]',' ',text)
            article_cleaned = article_cleaned + text
        article_cleaned = article_cleaned.lower()
        article_cleaned = article_cleaned.split()
        article_cleaned = [ps.stem(word) for word in article_cleaned if not word in set(stopwords.words('english'))]
        article_cleaned = ' '.join(article_cleaned)
        count = word_count(article_cleaned)
        sorted_x = sorted(count.items(), key=operator.itemgetter(1), reverse = True)
        real_estate_hits = hit_count(sorted_x, real_estate_keywords)
        retail_hits  = hit_count(sorted_x , retail_keywords)
        energy_hits = hit_count(sorted_x, energy_keywords)
        if fund_tags == 'energy' and energy_hits > 10:
            article_list.append(article)
        if fund_tags == 'real_estate' and real_estate_hits > 10:
            article_list.append(article)
    i=1
    print('Relevant Articles : ')
    for article in article_list:
        print(str(i) + '-> '+ article)
        i = i + 1